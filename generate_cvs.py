import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


def set_paragraph_spacing(paragraph, before=0, after=0, line=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line


def set_run_font(run, name='Outfit', size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=14, after=6)
    run = p.add_run(text.upper())
    set_run_font(run, name='Outfit', size=11, bold=True, color=(0x33, 0x33, 0x33))
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '333333')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_bullet(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=1, after=1, line=1.15)
    p.style = doc.styles['List Bullet']
    run = p.add_run(text)
    set_run_font(run, size=10.5)


def add_body_text(doc, text, size=10.5, bold=False, italic=False):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=2, after=2, line=1.15)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)


def add_experience_item(doc, title, org, date, bullets):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=1)
    run = p.add_run(title)
    set_run_font(run, size=11, bold=True)
    run = p.add_run(f'  \u2014  {org}')
    set_run_font(run, size=10.5, italic=True)

    p2 = doc.add_paragraph()
    set_paragraph_spacing(p2, before=0, after=2)
    run = p2.add_run(date)
    set_run_font(run, size=10, italic=True, color=(0x66, 0x66, 0x66))

    for bullet in bullets:
        add_bullet(doc, bullet)


def build_cv():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    style = doc.styles['Normal']
    style.font.name = 'Outfit'
    style.font.size = Pt(10.5)

    # --- Header ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=2)
    run = p.add_run('NDICHIA QUINCY FIEN')
    set_run_font(run, size=16, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=2)
    run = p.add_run('Full-Stack Developer \u2192 Cloud Security Engineer')
    set_run_font(run, size=12, color=(0x55, 0x55, 0x55))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=6)
    run = p.add_run('quincyfien99@gmail.com  |  +237 653 319 958  |  Bamenda, Cameroon')
    set_run_font(run, size=9.5, color=(0x66, 0x66, 0x66))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=8)
    run = p.add_run('github.com/quincyfien  |  linkedin.com/in/ndichia-quincy')
    set_run_font(run, size=9.5, color=(0x66, 0x66, 0x66))

    # --- Professional Summary ---
    add_section_heading(doc, 'Professional Summary')
    add_body_text(doc,
        'Full-stack developer with a physics-trained, first-principles approach to engineering, actively '
        'growing into a cloud security engineer. I write the technical documentation before I write the code, '
        'build web applications end-to-end, and apply secure coding practices at every step. Currently pursuing '
        'a Professional Master\u2019s degree in Cybersecurity, combining a deep understanding of how software is '
        'built with how it is attacked.',
        size=10.5,
    )

    # --- Technical Skills ---
    add_section_heading(doc, 'Technical Skills')
    skills_data = [
        ('Languages:', 'Python, JavaScript, TypeScript, SQL, HTML, CSS'),
        ('Frontend:', 'React, Next.js, Tailwind CSS'),
        ('Backend:', 'Django, Django REST Framework, FastAPI, Express.js, Node.js'),
        ('Databases & Infra:', 'PostgreSQL, MongoDB, Docker, Git & version control, REST APIs'),
        ('Security (Actively Building):', 'Network security, penetration testing (Nmap, Wireshark), Suricata IDS, IAM, incident response, cloud fundamentals (AWS), secure SDLC'),
    ]
    for label, content in skills_data:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=1, after=1, line=1.15)
        run = p.add_run(label + ' ')
        set_run_font(run, size=10.5, bold=True)
        run = p.add_run(content)
        set_run_font(run, size=10.5)

    # --- Education ---
    add_section_heading(doc, 'Education')
    add_bullet(doc, 'Professional Master\u2019s degree in Cybersecurity \u2014 University of Bamenda (2025\u2013Present)')
    add_bullet(doc, 'Bachelor of Science in Physics \u2014 University of Bamenda (2021\u20132025)')

    # --- Projects ---
    add_section_heading(doc, 'Projects')

    add_experience_item(doc,
        'Effica \u2014 Enterprise Printing Press Management Platform',
        'Full-Stack + Security  \u00b7  Django, PostgreSQL, Redis, Celery, Docker',
        '',
        [
            'Built a multi-branch SaaS platform handling the full lifecycle of a printing business: ordering, Mobile Money payments, production workflow, inventory, and cross-branch reporting.',
            'Designed a three-portal isolated architecture with hybrid RBAC + ABAC access control and branch-scoped row-level security.',
            'Implemented a custom WAF middleware, hash-chained tamper-proof audit logs, passwordless JWT authentication, and HMAC-SHA256 webhook verification for payments.',
            'Shipped with a CI/CD pipeline running 131 automated tests at an 80% coverage gate.',
        ]
    )
    add_experience_item(doc,
        'Circuit Forge \u2014 Custom PC E-Commerce Platform',
        'Full-Stack  \u00b7  Django, JavaScript, PostgreSQL, Docker',
        '',
        [
            'Built an e-commerce platform for custom PC building with a dynamic component compatibility checker.',
            'Designed a relational schema mapping socket types, power requirements, and dimensions to prevent invalid builds at checkout.',
            'Integrated real-time inventory tracking, multistep checkout, and an admin analytics dashboard.',
        ]
    )
    add_experience_item(doc,
        'SecureBank Incident Simulation \u2014 Red Team / Blue Team',
        'Security  \u00b7  Suricata IDS, Nmap, Wireshark, Linux',
        '',
        [
            'Conducted web exploitation and privilege escalation (Red Team), then hardened systems and deployed Suricata IDS (Blue Team).',
            'Engineered Suricata threshold rules that reduced false positives by 85% while catching stealthy Nmap scans.',
        ]
    )
    add_experience_item(doc,
        'Malware Classification System',
        'Security / ML  \u00b7  Python, Scikit-Learn, Pandas',
        '',
        [
            'Built a static feature extraction pipeline and compared SVM, Neural Network, and GAN-augmented classifiers for malware detection from PE headers.',
            'Applied SMOTE oversampling and hyperparameter tuning to boost recall on rare malware variants.',
        ]
    )
    add_experience_item(doc,
        'Replicated Key-Value Store',
        'Distributed Systems  \u00b7  Python, Sockets',
        '',
        [
            'Implemented a fault-tolerant master-replica key-value store with synchronous replication and heartbeat failure detection.',
            'Achieved read-your-writes consistency via client vector clocks and logical timestamps.',
        ]
    )

    # --- Core Competencies ---
    add_section_heading(doc, 'Core Competencies')
    competencies = [
        'Full-Stack Development \u2014 End-to-end delivery from database schema to user interface',
        'Secure Development \u2014 Threat modeling, secure coding, OWASP Top 10 mitigations',
        'Technical Writing \u2014 Documenting architecture and security requirements before implementation',
        'Analytical Problem Solving \u2014 Physics-trained first-principles reasoning and debugging',
        'Systems Thinking \u2014 Distributed architecture, consistency models, and failure handling',
    ]
    for comp in competencies:
        add_bullet(doc, comp)

    # --- Interests ---
    add_section_heading(doc, 'Interests')
    add_body_text(doc, 'Cloud Security  \u00b7  Secure SDLC  \u00b7  Distributed Systems  \u00b7  Open Source')

    path = os.path.join(OUTPUT_DIR, 'CV_Ndichia_Quincy.docx')
    doc.save(path)
    print(f'Saved: {path}')


if __name__ == '__main__':
    build_cv()
